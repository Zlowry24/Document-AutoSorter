from operator import truediv
from pathlib import Path  #imports the path class from pathlib python package
from docx import Document #imports Document class from PyDocx package

folder_path = Path(r"C:\Users\MyName\Desktop\Unsorted Documents") #points to folder, rstring prevents backspace from causing issues
files = folder_path.iterdir()  #Returns an iterator of path objects for each file/ folder in Unsorted documents (.iterdir is from the path class
for file in files: #loops through files variable(iterator of path objects) temp naming each one file
    print(file)

file_path = Path(r"C:\Users\lowry\Desktop\Unsorted Documents\FIAR_SOW_Document.docx") #points to file
doc = Document(file_path) #doc becomes a document object from the pydocx package, holding everything inside the docx file, in the provided location

subject_keywords = {'Audit' : ('FIAR','audit','internal controls','NGA', 'DHRA',
                               'Internal Controls', 'Obligation Monitoring', 'IUS', "DAR-Q"), #Dictionary containing subjects, and keywords that relate to them
                   'Budget' : ('budget',"PPBE","Programming",'BFA'),
                    'Finance & Accounting' : ('finance', 'accounting', 'GLO', 'DAR-Q','General Ledger'),
                    'Admin ' : ('admin', 'administrative', 'Executive Assistant'),
                    'Brain Health & Human Research' : ('brain health', 'health', 'research'),
                    'Contract Specialist Support' : ('contract specialist', 'contract support','contract', 'Acquisition'),
                    'Human Resources' : ('manpower augmentation', 'manpower', 'staffing needs'),
                    'Logistics & Doctrines' : ('logistics', 'doctrines','doctrine'),
                    'Security': ('security','security officer'),
                    'Strategic Planning Support': ('strategic planning support', 'strategic planning', 'strategic','strategic specialist'),



                    }


def determine_subject(file):
    scores = {"Audit": 0, 'Budget': 2, "Finance & Accounting":0, "Admin":0,
              "Brain Health & Human Research":0, "Contract Specialist Support":0,
              "Human Resources":0,"Logistics & Doctrines":0,"Security":0,
              "Strategic Planning Support":0, } #Dict of subject scores

    for paragraph in doc.paragraphs:  # loops through doc variable with .paragraphs attribute, which is a list of all the paragraphs of text (just paragraphs) in the document object
        print(paragraph.text) #prints paragraphs
        for subject_key, keywords in subject_keywords.items(): #loops through dictionary as a list using .items
            for keyword in keywords:  #checks (loops through) each individual keyword
                if keyword.lower() in paragraph.text.lower(): #checks if paragraph text matches any keywords during loop
                    scores[subject_key] += 1 #updates the score for the subjects in the score dictionary, Subject_key Var works bc the strings in both dictionaries are the same.

    best_subject = max(scores, key=scores.get) #Returns the subject(key) with the highest value
    best_value = scores[best_subject] #Returns the value of the best_subject Key
    best_matches = [] #empty list for the best subjects
    print("-"*100)
    print("List of the Secondary Subjects ")
    for score_key, score_value in scores.items(): #loops through scores dictionary making a list of keys and values
         if score_value == best_value:
             best_matches.append(score_key)
         elif score_value < best_value and score_value != 0:
             print(f"\n{score_key}: {score_value} matches")




    if len(best_matches) == 0:
        print("No matches!")
    elif len(best_matches) == 1:
        print("-"*100)
        print("\nBEST Subject Match:", best_matches[0], "With",best_value,"Matches")
        print("\n"+"-" * 100)
    else:
        print("\nTied for Best Subject Match Are", best_matches,"With",best_value,"Matches")









""" What this means, Go through each paragraph, and check all subjects(keys) for keyword matches in that paragraph, then print good job everytime theres a match"""




determine_subject(file_path)


