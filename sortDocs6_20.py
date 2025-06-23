import shutil
from pathlib import Path  #imports the path class from pathlib python package
from sys import prefix

from docx import Document #imports Document class from PyDocx package
from PyPDF2 import PdfReader



"""Checks if a txt file is being run, if not creates one with default dictionary info"""

def subject_keywords_txt(txtfile):
    if not txtfile.exists():
        with open(txtfile, "w", encoding="utf-8") as file:
            file.write("#Format: Subject: keyword1, keyword2, ...\n")
            file.write("Audit: FIAR, audit, internal controls, NGA, DHRA, Obligation Monitoring, IUS, DAR-Q\n")
            file.write("Budget: budget, PPBE, Programming, BFA\n")
            file.write("Admin: admin, administrative, Executive Assistant\n")
            file.write("Brain Health & Human Research: brain health, health, research\n")
            file.write("Contract Specialist Support: contract specialist, contract support, contract, Acquisition\n")
            file.write("Human Resources: manpower augmentation, manpower, staffing needs\n")
            file.write("Logistics & Doctrines: logistics, doctrines, doctrine\n")
            file.write("Security: security, security officer\n")
            file.write("Strategic Planning Support: strategic planning support, strategic planning, strategic, strategic specialist\n")
            file.write("Change Management: change management\n")
            file.write("Program Management Support: program management support, program management\n")
            file.write("Systems Engineering: systems engineering\n")
            #missing assesments, event planning, exercises, mgmt approach, req dev, realty and facilities


"""opens txt file, creates an empty dict, loops through each line, strips each line of spaces and newlines, ignores comments and empty lines, splits keys and values, and assigns values, splits values through commas, and makes a list of each, puts it in a tuple as the value for the corresponding subject."""

def subject_keywords_dict(txtfile):
    with open(txtfile, "r", encoding="utf-8") as file:
        subject_keywords_dict_two = {}
        for line in file:
            line = line.strip()
            if line.startswith("#") or line == "":
                continue
            subject, keywords_str = line.split(":",1) #splits the line at first colon into a list of index 0 and 1, and sets subject = to index 0 etc.
            subject = subject.strip()
            keywords_raw = keywords_str.split(",")
            cleaned_keywords = []
            for keyword in keywords_raw:
                cleaned = keyword.strip()
                if cleaned:
                    cleaned_keywords.append(cleaned)

            subject_keywords_dict_two[subject] = tuple(cleaned_keywords)
    return subject_keywords_dict_two




keyword_file = Path("subject_keywords.txt")
subject_keywords_txt(keyword_file)  # Ensures default exists
subject_keywords = subject_keywords_dict(keyword_file)





""" Dictionary containing subjects of business (keys), and keywords used to classify the subjects (Values) """



""" Stores the highest scoring subject with max(), and stores its value. Uses a for loop to add abs max scores to a list, and prints keys and values of subjects greater than 1, but not = 0.  Checks highest scores list length to determine if to print no matches, the best score, or all the best scores. """


def get_subject_results(scores, filename = None):
    if all(value == 0 for value in scores.values()): #checks if scores are all 0, before determining a best subject
        print(f"No keyword matches {filename}  cannot be classified.")
        return None

    best_subject = max(scores, key=scores.get)  # Returns the subject(key) with the highest value
    best_value = scores[best_subject]  # Returns the value of the best_subject Key
    best_matches = []  # empty list for the best subjects

    for score_key, score_value in scores.items():  # loops through scores dictionary making a list of keys and values
        if score_value == best_value:  # If a score value is equal to the max score,
            best_matches.append(score_key)  # add it to the best_matches list
        elif score_value < best_value and score_value != 0:  # If a score value is < max but >=1,
            print("-" * 100)
            print(f"List of the Secondary Subjects for {filename}:  ")
            print(f"\n{score_key}: {score_value} matches")  # print the subsequent score_key and its value



    if len(best_matches) == 1:
        print("-" * 100)
        print(f"\nBEST Subject Match for {filename}:", best_matches[0], "With", best_value, "Matches")
        print("\n" + "-" * 100)
        return best_subject

    else:
        print(f"\n{filename} has a tie for Best Subject Match between", best_matches, "With", best_value, "Matches")
        return best_subject




""" Function that uses nested for loops and if statements to check Paragraphs in a word document, for keywords within those paragraphs, and if a keyword is found in the text, it adds a point to the relating subject key  """
def determine_subject_docx(file_path):
    try:
        doc = Document(file_path)  # doc becomes a document object from the pydocx package, holding everything inside the docx file, in the provided location

    except Exception as e:
        print(f"\n Error opening {file_path.name}: {e}")
        return None  # No subject match for this file

    scores = {subject: 0 for subject in subject_keywords}  # Dict of subject scores created based off subjects, and set = 0

    for paragraph in doc.paragraphs:  # loops through doc variable with .paragraphs attribute, which is a list of all the paragraphs of text (just paragraphs) in the document object
        for subject_key, keywords in subject_keywords.items(): #loops through dictionary as a list using .items
            for keyword in keywords:  #checks (loops through) each individual keyword
                if keyword.lower() in paragraph.text.lower(): #checks if paragraph text matches any keywords during loop
                    count = paragraph.text.lower().count(keyword.lower())  #stores count of paragraphs text(lowercase) and within that counts keywords  (lowercase)
                    scores[subject_key] += count #Adds the count to the score value and sets it equal each time
                    for _ in range(count):
                        print(f"COUNT +1: Matched keyword '{keyword}' in: {paragraph.text}")

    for table in doc.tables: #the same as above but loops through each table, and its rows and cells.
        for row in table.rows:
            for cell in row.cells:
                for subject_key, keywords in subject_keywords.items():
                    for keyword in keywords:
                        if keyword.lower() in cell.text.lower():
                            count = cell.text.lower().count(keyword.lower())
                            scores[subject_key] += count
                            for _ in range(count):
                                print(f"COUNT +1: Matched keyword '{keyword}' in table cell: {cell.text}")

    for section in doc.sections: #same thing as above but it loops through each "section" which include headers and footers.
        for paragraph in section.header.paragraphs:
            for subject_key, keywords in subject_keywords.items():
                for keyword in keywords:
                    if keyword.lower() in paragraph.text.lower():
                        count = paragraph.text.lower().count(keyword.lower())
                        scores[subject_key] += count

    for paragraph in section.footer.paragraphs:
        for subject_key, keywords in subject_keywords.items():
            for keyword in keywords:
                if keyword.lower() in paragraph.text.lower():
                    count = paragraph.text.lower().count(keyword.lower())
                    scores[subject_key] += count

    best_subject = get_subject_results(scores,file_path.name) #keeps best_subject alive (it only exists in the get subject_results function)
    return best_subject



""" Function that uses nested for loops and if statements to check text lines in a PDF file for keywords within those text lines, and if a keyword is found in the text, it adds a point to the relating subject key  """

def determine_subject_pdf(file_path):
   pdf = PdfReader(file_path)

   scores = {subject:0 for subject in subject_keywords}# Dict of subject scores created based off subjects, and set = 0

   for page in pdf.pages: #same as logic as docx, but instead of paragraphs, it loops through the entire text (getting rid of lines in between).
       text = page.extract_text()
       if text:
           full_text = text.replace('\n', ' ')
           for subject_key, keywords in subject_keywords.items():
                for keyword in keywords:
                    if keyword.lower() in full_text.lower():
                        count = full_text.lower().count(keyword.lower())
                        scores[subject_key] += count
                        for _ in range(count):
                            print(f"COUNT +1: Matched keyword '{keyword}' in PDF page.")
   best_subject = get_subject_results(scores,file_path.name)
   return best_subject



def auto_sort(input_folder, output_folder):

    for file in input_folder.iterdir(): #loop through each content in the folder (itedir from pathlib)
        subject = None #prevents python crash if file is neither

        if file.suffix.lower() == ".docx":
            subject = determine_subject_docx(file)
        elif file.suffix.lower() == ".pdf": #elif bc they're mutually exclusive, if statements checks both regardless if one already matched.
            subject = determine_subject_pdf(file)

        if subject:
                subject_folder = output_folder / subject  # assigns value of subject folder to folder directory with /(whatever subject) added on. example "C:/Users/lowry/Desktop/Sorted Documents/Budget"

                subject_folder.mkdir(exist_ok=True)

                destination = subject_folder / file.name #.name is built in to path gives file name with no folder path as a string

                shutil.copy(file, destination)











""" Menu that runs until exited, decides what function to run and allows for input of file paths"""
def main():
    while True:
        print("\nMenu")
        print("1. Classify Word Documents")
        print("2. Classify PDF")
        print("3. Auto Sort")
        print("4. Import Subject Dictionary")
        print("5. Exit")
        input_choice = input("Enter your choice: ")
        if input_choice == "1":
            path_str = input("Enter the path of your document: ").strip('"')  # asking user for path, and negates parenthesis
            file_path = Path(path_str)  # turning the string into a path object
            determine_subject_docx(file_path)
        if input_choice == "2":
            pdf_path_str = input("Enter the path of your PDF: ").strip('"')
            file_path = Path(pdf_path_str)
            determine_subject_pdf(file_path)
        if input_choice == "3":
            input_path_str = input("Enter the folder path of your documents: ").strip('"')
            output_path_str = input("Enter the destination folder: ").strip('"')
            input_folder = Path(input_path_str)
            output_folder = Path(output_path_str)
            auto_sort(input_folder, output_folder)

        if input_choice == "4":
            path_str_txt = input("Enter the path of your txt file: ").strip('"')
            keyword_file_user = Path(path_str_txt)
            global subject_keywords #replace global subject_keywords dict with this one
            subject_keywords_txt(keyword_file_user)
            subject_keywords = subject_keywords_dict(keyword_file_user)
            print("New subject keywords loaded:", subject_keywords)

        if input_choice == "5":
            print("Goodbye!")
            break







if __name__ == "__main__":
    main()