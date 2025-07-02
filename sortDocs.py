import shutil
import re
from pathlib import Path  #imports the path class from pathlib python package
from docx import Document  #imports Document class from PyDocx package
from PyPDF2 import PdfReader


Debug = True #sets Debug mode to false
min_keyword_threshold = 4
overwrite_value = False


"""Checks if a txt file is being run, if not creates one with default dictionary info"""
def subject_keywords_txt(txtfile, overwrite = False):
    if not txtfile.exists() or overwrite:
        with open(txtfile, "w", encoding="utf-8") as file:
            file.write("#Format: Subject: keyword1, keyword2, ...\n")
            file.write("Audit: FIAR, audit, audit readiness, internal controls, NGA, DHRA, Obligation Monitoring, IUS, Internal Use Software, IT Hardware, KSD, Key Supporting Documentation, Package, NFR, Notice of Findings and Recommendations, CAP, Corrective Action Plan, MAP, SOA, Statement of Assurance\n")
            file.write("Finance & Accounting: DAR-Q, FPS, DIA, USCG, Accounting, Finance, Reconciliation, General Ledger, Journal Entry, GL, JE, PP&E, Property Plant and Equipment, Financial Planning and Analysis, FPA\n")
            file.write("Budget: budget, PPBE, Programming, BFA, Execution, planning, POM, FYDP, DEAMS, FEM, PPBES, PPBES-MIS, MIPR, FORM 9, MORD, PROJECT ORDER, 4009, GPC, SUPPORT AGREEMENTS, GFEBS, DAI, CCaR, GAFS-BQ, ABSS, DFAS, BES, PB, AUTHORITIES, PTEO, LOA\n")
            file.write("Admin: admin, administrative, Executive Assistant, EA\n")
            file.write("Brain Health & Human Research: brain health, health, research, human research, HRPO, PEO-ST, Surgeon, Science and Technology, Science & Technology, S&T, CASH, ANAM, TBI, Brain Injury, HRPP, BUMED\n")
            file.write("Contract Specialist Support: contract specialist, contract support, contract, Acquisition, contract specialist support, Procurement, Pre-award, post-award, contract closeout, contract documentation, PEO-K\n")
            file.write("Human Resources: manpower augmentation, manpower, staffing needs, HR, HR generalist, HR Specialist, Human Resources\n")
            file.write("Logistics & Doctrines: logistics, doctrines, doctrine, J4, Directives\n")
            file.write("Security: security, security officer, security specialist, WHS, SSO, J2-SSO\n")
            file.write( "Strategic Planning Support: strategic planning support, strategic planning, strategic, strategic specialist, strategic planner, POAM\n")
            file.write("Change Management: change management, CRIT, HOSC, Deloitte\n")
            file.write("Program Management Support: program management support, program management, PEO-M\n")
            file.write("Systems Engineering: systems engineering, engineer, SDA, DCGS\n")
            file.write("Management Approach: Transition Plan, Transition, Recruiting, Hiring, Training, Staff, Staff training\n")
            file.write("Assessments: N2C3, LCSP, JAF\n")
            file.write("Event Planning: event planning, event, ISOF\n")
            file.write("Exercises: Exercises, Range, CMRC\n")
            file.write("Requirements Development: JCIDS, SOFCIDS, Requirements, JRIMS, CDD, ICD, DCR\n")
            file.write("Realty: Realty, Real estate\n")
            file.write("Facilities: Facilities\n")



            #missing assessments, event planning, exercises, req dev, realty and facilities


"""opens txt file, creates an empty dict, loops through each line, strips each line of spaces and newlines, ignores comments and empty lines, splits keys and values, and assigns values, splits values through commas, and makes a list of each, puts it in a tuple as the value for the corresponding subject."""


def subject_keywords_dict(txtfile):
    with open(txtfile, "r", encoding="utf-8") as file:
        subject_keywords_dict_two = {}
        for line in file:
            line = line.strip()
            if line.startswith("#") or line == "":
                continue
            subject, keywords_str = line.split(":",
                                               1)  #splits the line at first colon into a list of index 0 and 1, and sets subject = to index 0 etc.
            subject = subject.strip()
            keywords_raw = keywords_str.split(",")
            cleaned_keywords = []
            for keyword in keywords_raw:
                cleaned = keyword.strip()
                if cleaned:
                    cleaned_keywords.append(cleaned)

            subject_keywords_dict_two[subject] = tuple(cleaned_keywords)
    return subject_keywords_dict_two




""" Stores the highest scoring subject with max(), and stores its value. Uses a for loop to add abs max scores to a list, and prints keys and values of subjects greater than 1, but not = 0.  Checks highest scores list length to determine if to print no matches, the best score, or all the best scores. """


def get_subject_results(scores, filename=None):
    if all(value == 0 for value in scores.values()):  #checks if scores are all 0, before determining best subject
        print(f"---------------- {filename} ----------------")
        print("No Keyword Matches. Cannot Be classified")
        print("--------------------------------------------------------\n")
        return None

    best_subject = max(scores, key=scores.get)  # Returns the subject(key) with the highest value
    best_value = scores[best_subject]  # Returns the value of the best_subject Key
    best_matches = []  # empty list for the best subjects
    qualified_secondary_matches = []
    secondary_matches = []
    if best_value < min_keyword_threshold:
        print(f"---------------- {filename} ----------------")
        print(f"Insufficient Keyword Matches.(Fewer than {min_keyword_threshold}  Cannot Be Classified.)")
        print("--------------------------------------------------------\n")
        return "Unclassified"

    for score_key, score_value in scores.items():  # loops through scores dictionary making a list of keys and values
        if score_value == best_value:  # If a score value is equal to the max score,
            best_matches.append(score_key)  # add it to the best_matches list
        if best_value > score_value >= 1:  # If a score value is < max but >=1,
            secondary_matches.append(score_key)
        if best_value > score_value >= min_keyword_threshold: qualified_secondary_matches.append(score_key)

    if len(secondary_matches) >= 1:
        print(f"\n---------------- {filename} ----------------")
        print("Secondary Subjects:  ")

        for subject in secondary_matches:
            score_value = scores[subject]
            print(f" - {subject}: {score_value} match{'es'if score_value !=1 else ''}")



    if len(best_matches) == 1:

        print(f"\nBEST Subject Match: {best_matches[0]} With ({best_value} Matches)")
        print("--------------------------------------------------------\n")
        return best_subject,qualified_secondary_matches

    else:
        print(f"\n{filename} has a tie for Best Subject Match between", best_matches, "With", best_value, "Matches")
        return best_subject,qualified_secondary_matches


""" Loops thorugh paragraphs, tables, headers, and footers adding them all to full text. Then loops through a dictionary as a list, sorts the keywords from longest to shortest, checks if a keyword is in the text, counts how many times it appears, adds the count to the score, and erases the keyword from the text"""

def determine_subject_docx(file_path):
    name = file_path.name
    full_text = ""
    try:
        doc = Document(
            file_path)  # doc becomes a document object from the pydocx package, holding everything inside the docx file, in the provided location

    except Exception as e:
        print(f"\n Error opening {file_path.name}: {e}")
        return None  # No subject match for this file

    scores = {subject: 0 for subject in
              subject_keywords}  # Dict of subject scores created based off subjects, and set = 0

    for paragraph in doc.paragraphs:  # loops through doc variable with .paragraphs attribute, which is a list of all the paragraphs of text (just paragraphs) in the document object
        full_text += paragraph.text + " "

    for table in doc.tables:  #the same as above but loops through each table, and its rows and cells.
        for row in table.rows:
            for cell in row.cells:
                full_text += cell.text + " "

    for section in doc.sections:  #same thing as above but it loops through each "section" which include headers and footers.
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:  #runs are headings, bold, and italics
                full_text += run.text + " "


    for paragraph in section.footer.paragraphs:
        for run in paragraph.runs:
            full_text += paragraph.text + " "
    text_to_search = full_text.lower()

    for subject_key, keywords in subject_keywords.items():  # loops through dictionary as a list using .items
        length_keywords = sorted(keywords, key=lambda k:-len(k)) #sorted function sorts list from smallest to largest, this sorts keywords by assigning negative value to each letter, making the longest word the smallest by value.
        for keyword in length_keywords:  # checks (loops through) each individual keyword sorted form longest to shortest
            pattern =r"\b" + re.escape(keyword.lower()) + r"\b" #uses re module to add a boundary after each keyword, so it only looks for (example) "EA" not ea in every word
            matches = re.findall(pattern, text_to_search) #searches entire text for all non-overlapping matches of a pattern and returns a list of strings
            if matches:  # checks if text matches any keywords during loop
                count = len(matches)  # counts matches in matches list
                scores[subject_key] += count  # Adds the count to the score value and sets it equal each time
                if Debug:
                    for _ in range(count):
                        print(f"COUNT +1: Matched keyword '{keyword}' in Document")
                text_to_search = re.sub(pattern," " *len(keyword.lower()), text_to_search) #removes occurance of keyword after counting it to avoid double counting, replacing it w/spaces = to the keyword length
    best_subject = get_subject_results(scores, file_path.name)  #keeps best_subject alive (it only exists in the get subject_results function)
    return best_subject, name


""" Function that uses nested for loops and if statements to check text lines in a PDF file for keywords within those text lines, and if a keyword is found in the text, it adds a point to the relating subject key  """


def determine_subject_pdf(file_path):
    pdf = PdfReader(file_path)
    name = file_path.name
    full_text = ""

    scores = {subject: 0 for subject in subject_keywords}  # Dict of subject scores created based off subjects, and set = 0

    for page in pdf.pages:  #same as logic as docx, but instead of paragraphs, it loops through the entire text (getting rid of lines in between).
        text = page.extract_text()
        if text:
            full_text = text.replace('\n', ' ')

            text_to_search = full_text.lower() #make a variable for full text so when it uses updated version not original

            for subject_key, keywords in subject_keywords.items():
                length_keywords = sorted(keywords, key=lambda k:-len(k))
                for keyword in length_keywords:
                    pattern = r"\b" + re.escape(keyword.lower()) + r"\b"  # uses re module to add a boundary after each keyword, so it only looks for (example) "EA" not ea in every word
                    matches = re.findall(pattern, text_to_search)  # searches entire text for all non-overlapping matches of a pattern and returns a list of strings
                    if matches:  # checks if text matches any keywords during loop
                        count = len(matches)  # counts matches in matches list
                        scores[subject_key] += count  # Adds the count to the score value and sets it equal each time

                        if Debug:
                            for _ in range(count):
                                print(f"COUNT +1: Matched keyword '{keyword}' in PDF page.")
                        text_to_search = re.sub(pattern, " " * len(keyword.lower()),text_to_search)  # removes occurance of keyword after counting it to avoid double counting, replacing it w/spaces = to the keyword length
    best_subject = get_subject_results(scores, file_path.name)
    return best_subject, name


def auto_sort(input_folder, output_folder,doc_type = None):
    log = {}
    secondary_log = {}
    total_files = 0
    for file in input_folder.iterdir():  #loop through each content in the folder (itedir from pathlib)
        if file.is_dir():
            continue #skip folders so, program doesn't try to copy them
        subject = None  #prevents python crash if file is neither

        if file.suffix.lower() == ".docx":
            subject, name = determine_subject_docx(file)
        elif file.suffix.lower() == ".pdf":  #elif bc they're mutually exclusive, if statements checks both regardless if one already matched.
            subject, name = determine_subject_pdf(file)
        if subject is None:
            subject = "Unclassified"

        if isinstance(subject, tuple): #This checks if the subject object is a tuple, if it is unpacked subject as best subject, and secondary matches normally
            best_subject, secondary_matches = subject
        else:
            best_subject, secondary_matches = subject, [] #if subject object isn't a tuple (just a str) make an empty list to pass through to avoid value error (not enough values to unpack)

        folder_path = output_folder / best_subject  # assigns value of subject folder to folder directory with /(whatever subject) added on. example "C:/Users/lowry/Desktop/Sorted Documents/Budget"

        if doc_type:
            folder_path = folder_path / doc_type  # assigns value of subject folder to folder directory with /(whatever subject) added on. example "C:/Users/lowry/Desktop/Sorted Documents/Budget"

        folder_path.mkdir(parents = True, exist_ok=True) #if any folders in path don't exist, make them all.


        shutil.copy(file,folder_path / file.name)

        if best_subject not in log:
             log[best_subject] = []  #make a list to store file names inside dictionary
        log[best_subject].append(file.name)

        total_files += 1

        for secondary in secondary_matches:
            sec_path = output_folder / secondary
            if doc_type:
                sec_path = sec_path / doc_type

            sec_path.mkdir(parents = True, exist_ok=True)

            shutil.copy(file, sec_path / file.name)

            if secondary not in secondary_log: #logs files that got sorted
                secondary_log[secondary] = []
            secondary_log[secondary].append(file.name)

            total_files += 1

    combined_log = {}
    for subject, files in log.items(): #loop through subjects and "value pair" of lists of file names
        if subject not in combined_log:
            combined_log[subject] = [] #make an empty list in combined log for where file names go
        combined_log[subject].extend(files) #transfer data from log(looping through) to combined_log

    for subject, files in secondary_log.items(): #loop through subjects and "value pair" of lists of file names
        if subject not in combined_log:
            combined_log[subject] = [] #make an empty list in combined log for where file names go
        combined_log[subject].extend(files) #transfer data from log(looping through) to combined_log

    print(f"\nTotal Files Sorted: {total_files}")
    print("\nWhere Everything Was Sorted:")
    for subject, files in combined_log.items():
        print(f"\n{subject}: {len(files)} file(s)")
        for filename in files:
            print(f"  - {filename}")




""" Menu that runs until exited, decides what function to run and allows for input of file paths"""


def main():
    global subject_keywords
    global overwrite_value
    global Debug
    keyword_file = Path("subject_keywords.txt")


    while True:
        subject_keywords_txt(keyword_file, overwrite=overwrite_value)  # Ensures default exists
        subject_keywords = subject_keywords_dict(keyword_file)
        if Debug:
            print("\n[DEBUG] Loaded keywords:")
            for subject, keywords in subject_keywords.items():
                print(f"{subject}: {keywords}")
        print("\nMenu")
        print("1. Classify Word Document")
        print("2. Classify PDF")
        print("3. Auto Sort (Subject)")
        print("4. Auto Sort (Subject + Type)")
        print("5. System Settings ")
        print("6. Exit")

        print("\nEnter back at any point if wrong selection is chosen.")
        input_choice = input("Enter your choice: ")
        if input_choice not in {"1", "2", "3", "4", "5","6"}:
            print("Invalid choice. Please try again.")
            continue
        if input_choice == "1":
            path_str = input("Enter the path of your document: ").strip( '"')  # asking user for path, and negates parenthesis
            if path_str.strip().lower() == "back": continue
            file_path = Path(path_str)  # turning the string into a path object
            determine_subject_docx(file_path)
        if input_choice == "2":
            pdf_path_str = input("Enter the path of your PDF: ").strip('"')
            if pdf_path_str.strip().lower() == "back": continue
            file_path = Path(pdf_path_str)
            determine_subject_pdf(file_path)
        if input_choice == "3":
            input_path_str = input("Enter the folder path of your documents: ").strip('"')
            if input_path_str.strip().lower() == "back": continue
            output_path_str = input("Enter the destination folder: ").strip('"')
            if output_path_str.strip().lower() == "back": continue
            input_folder = Path(input_path_str)
            output_folder = Path(output_path_str)
            auto_sort(input_folder, output_folder)

        if input_choice == "4":
            print("\nNote:\n  The document type you enter will be used exactly as-is to create a folder.")
            print(
                "To sort into an existing folder (like 'RFP' or 'SOW'), you must type the folder name exactly, including: ")
            print("- Correct capitalization (e.g., 'RFP' ≠ 'rfp')")
            print("- No extra spaces")
            print("- No added letters (e.g., 'RFPs' will create a new folder separate from 'RFP')")
            print("\n Even small changes will create a new folder.")

            input_type = input("\nEnter the type of Documents: ")
            if input_type.strip().lower() == "back": continue
            input_path_str = input("Enter the folder path of your documents: ").strip('"')
            if input_path_str.strip().lower() == "back": continue
            output_path_str = input("Enter the destination folder: ").strip('"')
            if output_path_str.strip().lower() == "back": continue
            input_folder = Path(input_path_str)
            output_folder = Path(output_path_str)
            auto_sort(input_folder, output_folder, input_type)


        if input_choice == "5": #New Menu for Customizations
            global min_keyword_threshold #declaring to python that im accessing global variables


            while True:
                print("\nSettings Menu")
                print(f"1. Set Minimum Keyword Threshold (current: {min_keyword_threshold}) ")
                print("2. Import New Subject Dictionary")
                print("3. Toggle Debug Mode")
                print("4. Toggle Update Default Dictionary File")
                print("5. Back to Main Menu")
                settings_choice = input("Enter your choice: ")
                if settings_choice.lower() not in {"1", "2", "3","4","5"}:
                    print("Invalid choice. Please try again.")
                    continue
                if settings_choice == "1": #Allows user to change the threshold in which a Document is Classified
                    min_keyword_threshold = int(input("Enter the minimum keyword threshold: "))
                if settings_choice == "2": # Allows a user to import a new subject dict for customized sorting
                    path_str_txt = input("Enter the path of your txt file: ").strip('"')
                    if path_str_txt.strip().lower() == "back": continue
                    keyword_file_user = Path(path_str_txt)
                    subject_keywords_txt(keyword_file_user) # replace global subject_keywords dict with this one
                    subject_keywords = subject_keywords_dict(keyword_file_user)
                    print("New subject keywords loaded:", subject_keywords)

                if settings_choice == "3":
                    if not Debug:
                        Debug = True
                        print("\nDebug mode ON")
                    elif Debug:
                        Debug = False
                        print("\nDebug mode OFF")
                if settings_choice == "4":
                    if not overwrite_value:
                        overwrite_value = True
                        print("\nDictionary Overwrite mode ON. The default dictionary will be Re-Updated each time the program runs.")
                    elif overwrite_value:
                        overwrite_value = False
                        print("\nDictionary Overwrite mode OFF, The default dictionary will be NOT Re-Updated each time the program runs, Status: Able to Import Custom Dictionary")
                if settings_choice == "5":
                    break

        if input_choice == "6":
            print("Goodbye!")
            break


if __name__ == "__main__":
    main()
